from __future__ import annotations
from typing import List, Tuple
import io
import os
import subprocess
import tempfile

def _safe_decode(b: bytes) -> str:
    try:
        return b.decode("utf-8")
    except Exception:
        try:
            return b.decode("latin-1")
        except Exception:
            return ""

def extract_text_from_upload(filename: str, data: bytes, cfg=None) -> Tuple[str, List[str]]:
    """Return (text, warnings). Local extraction only (no OCR).

    Supported:
      - .txt/.md/.log/.csv (decoded as text)
      - .docx (python-docx)
      - .xlsx/.xlsm (openpyxl)
      - .pdf (best-effort via pypdf if installed)

    Legacy formats (.doc/.xls) are accepted and will be converted locally via
    LibreOffice (soffice) if available.
    """
    warnings: List[str] = []
    name = (filename or "").lower()

    if name.endswith((".txt", ".md", ".log", ".csv")):
        return _safe_decode(data), warnings

    if name.endswith(".docx"):
        try:
            from docx import Document  # type: ignore
            doc = Document(io.BytesIO(data))
            parts: List[str] = []
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    parts.append(t)
            for table in doc.tables:
                for row in table.rows:
                    cells = [(c.text or "").strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts), warnings
        except Exception as e:
            warnings.append(f"DOCX parse failed: {e}")
            return "", warnings

    if name.endswith((".xlsx", ".xlsm")):
        try:
            from openpyxl import load_workbook  # type: ignore
            wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
            out_lines: List[str] = []
            for sheet in wb.worksheets:
                out_lines.append(f"## Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    if row and any(v is not None and str(v).strip() != "" for v in row):
                        out_lines.append(", ".join("" if v is None else str(v) for v in row))
            return "\n".join(out_lines), warnings
        except Exception as e:
            warnings.append(f"XLSX parse failed: {e}")
            return "", warnings

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception as e:
            warnings.append(f"PDF parsing requires pypdf. Install it in your venv. ({e})")
            return "", warnings

        try:
            reader = PdfReader(io.BytesIO(data))
            parts: List[str] = []
            for i, page in enumerate(reader.pages):
                try:
                    t = page.extract_text() or ""
                except Exception:
                    t = ""
                t = t.strip()
                if t:
                    parts.append(f"## Page {i+1}\n{t}")
            if not parts:
                warnings.append("PDF had no extractable text (may be scanned). Consider OCR later.")
            return "\n".join(parts), warnings
        except Exception as e:
            warnings.append(f"PDF parse failed: {e}")
            return "", warnings

    if name.endswith((".doc", ".xls")):
        if not bool(getattr(cfg, "allow_legacy_office_conversion", False)):
            warnings.append(
                f"Legacy Office conversion is disabled for safety: {filename}. "
                "Please save/export as .docx/.xlsx and upload again, or enable the setting if you trust the input."
            )
            return "", warnings

        # Best-effort local conversion using LibreOffice (soffice). This keeps
        # BYO-SecAI production-friendly without requiring cloud/OCR.
        try:
            ext = ".doc" if name.endswith(".doc") else ".xls"
            with tempfile.TemporaryDirectory(prefix="byo_secai_ingest_") as td:
                src_path = os.path.join(td, f"input{ext}")
                with open(src_path, "wb") as f:
                    f.write(data)

                # Convert to the modern equivalent so our native parsers can handle it.
                # - .doc  -> .docx
                # - .xls  -> .xlsx
                target = "docx" if ext == ".doc" else "xlsx"
                cmd = ["soffice", "--headless", "--nologo", "--nolockcheck", "--convert-to", target, "--outdir", td, src_path]
                proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=45)
                if proc.returncode != 0:
                    warnings.append(
                        f"Legacy Office conversion failed for {filename} (soffice rc={proc.returncode}). "
                        f"stderr: {(proc.stderr or b'').decode('utf-8', errors='ignore')[:300]}"
                    )
                    return "", warnings

                # Find converted output
                out_path = os.path.join(td, f"input.{target}")
                if not os.path.exists(out_path):
                    # LibreOffice sometimes changes filenames; pick the first matching target ext.
                    for fn in os.listdir(td):
                        if fn.lower().endswith(f".{target}"):
                            out_path = os.path.join(td, fn)
                            break

                if not os.path.exists(out_path):
                    warnings.append(f"Legacy Office conversion produced no .{target} output for {filename}.")
                    return "", warnings

                with open(out_path, "rb") as f:
                    converted = f.read()

            # Re-enter parsing with the converted payload.
            converted_name = filename + (".docx" if ext == ".doc" else ".xlsx")
            return extract_text_from_upload(converted_name, converted, cfg=cfg)
        except FileNotFoundError:
            warnings.append(
                f"Legacy Office format not supported for parsing: {filename}. LibreOffice (soffice) not found. "
                "Please save as .docx/.xlsx."
            )
            return "", warnings
        except Exception as e:
            warnings.append(f"Legacy Office conversion failed for {filename}: {e}")
            return "", warnings

    warnings.append(f"Unsupported file type: {filename}")
    return "", warnings
